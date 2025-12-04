import os

import pandas as pd
from docx import Document

import utils  # ログ出力とエラーハンドリング用


def create_sample_data():
    """
    サンプルのExcelデータとWordテンプレートを作成します。
    data/, templates/, output/ ディレクトリに必要なファイルを生成します。
    """
    # 処理開始ログ
    utils.log_start("create_sample_data")

    try:
        # ディレクトリの作成（セキュリティ: 存在確認とエラーハンドリング）
        directories = ["data", "templates", "output"]
        for directory in directories:
            try:
                os.makedirs(directory, exist_ok=True)
                utils.log_start(f"ディレクトリ作成: {directory}")
            except OSError as e:
                utils.log_error(
                    f"ディレクトリ '{directory}' の作成に失敗: {e}"
                )
                raise

        # 1. サンプルExcelデータの作成
        # 初心者向け解説: 辞書形式でデータを定義し、DataFrameに変換します
        data = {
            "property_name": [
                "グランドハイツ東京",
                "サニーサイド横浜",
                "リバーサイド大阪",
            ],
            "address": [
                "東京都千代田区1-1-1",
                "神奈川県横浜市西区2-2-2",
                "大阪府大阪市北区3-3-3",
            ],
            "amount": [100000, 150000, 80000],
        }

        # データフレームに変換
        df = pd.DataFrame(data)
        excel_path = "data/contract_data.xlsx"

        # Excelファイルへの書き込み（外部I/Oなのでエラーハンドリング）
        try:
            df.to_excel(excel_path, index=False)
            print(f"成功: Excelデータを作成しました: {excel_path}")
        except Exception as e:
            utils.log_error(f"Excelファイルの作成に失敗: {excel_path}")
            utils.handle_error(e)

        # 2. サンプルWordテンプレートの作成
        try:
            doc = Document()
            doc.add_heading("賃貸借契約書", 0)

            doc.add_paragraph(
                "貸主（以下「甲」という）と借主（以下「乙」という）は、"
                "以下の通り賃貸借契約を締結する。"
            )

            # 第1条（物件の表示）
            doc.add_heading("第1条（物件の表示）", level=1)
            p = doc.add_paragraph()
            p.add_run("物件名: ").bold = True
            p.add_run("{{property_name}}")

            p = doc.add_paragraph()
            p.add_run("住所: ").bold = True
            p.add_run("{{address}}")

            # 第2条（賃料）
            doc.add_heading("第2条（賃料）", level=1)
            p = doc.add_paragraph()
            p.add_run("賃料: 金 ").bold = True
            p.add_run("{{amount}}")
            p.add_run(" 円")

            # Wordファイルへの保存（外部I/Oなのでエラーハンドリング）
            template_path = "templates/contract_template.docx"
            doc.save(template_path)
            print(f"成功: Wordテンプレートを作成しました: {template_path}")

        except Exception as e:
            utils.log_error(f"Wordテンプレートの作成に失敗: {template_path}")
            utils.handle_error(e)

    except Exception as e:
        # 予期しないエラーが発生した場合
        utils.handle_error(e)

    # 処理終了ログ
    utils.log_end("create_sample_data")


if __name__ == "__main__":
    try:
        create_sample_data()
        print("サンプルアセットの作成が完了しました。")
    except Exception as e:
        utils.log_error(f"サンプルアセット作成中にエラーが発生: {e}")
        utils.handle_error(e)
